"""
SUTRIX Thesis DOCX Builder
Compiles the modular chapter data into a publication-grade Word Document (DOCX).
Implements professional academic styles, table formats, figure embeddings, and headers/footers.
Integrates Vancouver style citations and semantic duplicate checks.
"""

import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

# Import references database
from thesis_compiler.chapters.references import REFERENCES

# Keep track of duplicate paragraphs removed globally
DUPLICATES_REMOVED_COUNT = 0
GLOBAL_CITATION_MAP = {}

def set_cell_background(cell, color_hex):
    """Sets background color for a table cell."""
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def set_table_borders(table):
    """Applies a clean, publication-grade academic border style (no vertical borders)."""
    tblPr = table._tbl.tblPr
    borders_xml = f'''<w:tblBorders {nsdecls("w")}>
        <w:top w:val="single" w:sz="6" w:space="0" w:color="CCCCCC"/>
        <w:bottom w:val="single" w:sz="12" w:space="0" w:color="1B365D"/>
        <w:insideH w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>
        <w:left w:val="none"/>
        <w:right w:val="none"/>
        <w:insideV w:val="none"/>
    </w:tblBorders>'''
    tblPr.append(parse_xml(borders_xml))

def add_page_number(run):
    """Inserts a dynamic PAGE field code for page numbering."""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    r = run._r
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)

def add_toc(paragraph):
    """Inserts a dynamic Table of Contents field code."""
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    r = run._r
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)

def apply_text_formatting(paragraph, font_name="Times New Roman", size_pt=11, bold=False, italic=False, color_rgb=(0,0,0)):
    """Applies standard font formatting to a paragraph."""
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(6)
    
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(size_pt)
        run.bold = bold
        run.italic = italic
        run.font.color.rgb = RGBColor(*color_rgb)

def preprocess_docx_chapters(chapters, threshold=0.85):
    """Preprocesses chapters: resolves Vancouver citations and performs semantic deduplication."""
    global DUPLICATES_REMOVED_COUNT, GLOBAL_CITATION_MAP
    citation_map = {}
    citation_counter = 1
    duplicate_count = 0
    
    def replace_citation(match):
        nonlocal citation_counter
        bracket_content = match.group(1)
        keys = [k.strip() for k in bracket_content.split(",")]
        valid_keys = [k for k in keys if k in REFERENCES]
        
        if not valid_keys:
            return match.group(0)
            
        numbers = []
        for key in valid_keys:
            if key not in citation_map:
                citation_map[key] = citation_counter
                citation_counter += 1
            numbers.append(str(citation_map[key]))
            
        return "[" + ", ".join(numbers) + "]"
        
    for chap in chapters:
        preprocessed_content = []
        chapter_text_history = []
        
        for item in chap.CONTENT:
            tag = item[0]
            if tag in ["p", "h1", "h2", "h3"]:
                val = item[1]
                
                # Apply Vancouver citation replacement
                new_val = re.sub(r'\[([a-zA-Z0-9_, ]+)\]', replace_citation, val)
                
                # Semantic Deduplication check (p tag only)
                if tag == "p":
                    words = new_val.split()
                    if len(words) >= 10:
                        # 1. Exact hash deduplication
                        normalized_p = re.sub(r'\s+', ' ', new_val).strip()
                        p_hash = hash(normalized_p)
                        is_exact_dup = False
                        for prev_p in chapter_text_history:
                            if hash(re.sub(r'\s+', ' ', prev_p).strip()) == p_hash:
                                is_exact_dup = True
                                break
                        if is_exact_dup:
                            duplicate_count += 1
                            continue
                            
                        # 2. Semantic check via TF-IDF Cosine Similarity
                        if chapter_text_history:
                            try:
                                vectorizer = TfidfVectorizer(analyzer='char_wb', ngram_range=(3, 5)).fit_transform(chapter_text_history + [new_val])
                                vectors = vectorizer.toarray()
                                p_vector = vectors[-1].reshape(1, -1)
                                prev_vectors = vectors[:-1]
                                similarities = cosine_similarity(p_vector, prev_vectors)[0]
                                if max(similarities) > threshold:
                                    duplicate_count += 1
                                    continue # Discard duplicate paragraph
                            except Exception:
                                pass
                        chapter_text_history.append(new_val)
                        
                preprocessed_content.append((tag, new_val))
                
            elif tag == "list":
                bullets = item[1]
                new_bullets = [re.sub(r'\[([a-zA-Z0-9_, ]+)\]', replace_citation, b) for b in bullets]
                preprocessed_content.append(("list", new_bullets))
                
            elif tag == "table":
                headers, rows, caption = item[1], item[2], item[3]
                new_caption = re.sub(r'\[([a-zA-Z0-9_, ]+)\]', replace_citation, caption)
                preprocessed_content.append(("table", headers, rows, new_caption))
                
            elif tag == "figure":
                fig_name, caption = item[1], item[2]
                new_caption = re.sub(r'\[([a-zA-Z0-9_, ]+)\]', replace_citation, caption)
                preprocessed_content.append(("figure", fig_name, new_caption))
                
            else:
                preprocessed_content.append(item)
                
        # Overwrite chapter content with preprocessed content
        chap.CONTENT = preprocessed_content
        
    DUPLICATES_REMOVED_COUNT = duplicate_count
    GLOBAL_CITATION_MAP = citation_map
    return citation_map, duplicate_count

def build_docx(output_path, figures_dir, chapters_list):
    """Compiles the chapters list into the styled DOCX dissertation manuscript."""
    print("Pre-processing chapters (Deduplication & Citation Resolution)...")
    citation_map, dup_removed = preprocess_docx_chapters(chapters_list, threshold=0.85)
    print(f"  Semantic duplicates removed: {dup_removed}")
    print(f"  Unique references mapped: {len(citation_map)}")
    
    print("Initializing Word Document Compiler...")
    doc = Document()
    
    # ─── Margins Setup ────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.different_first_page_header_footer = True
        
        # Configure Footer Page Numbering (Arabic)
        footer = section.footer
        f_p = footer.paragraphs[0]
        f_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        f_run = f_p.add_run("Page ")
        add_page_number(f_run)
        apply_text_formatting(f_p, size_pt=9, color_rgb=(100, 100, 100))
        
        # Configure Header running title
        header = section.header
        h_p = header.paragraphs[0]
        h_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h_p.text = f"SUTRIX: A Regulatory-Grade Scientific Data Orchestration Platform"
        apply_text_formatting(h_p, size_pt=8.5, color_rgb=(100, 100, 100), italic=True)

    # ─── Title Page ───────────────────────────────────────────────────────────
    print("Compiling Cover Page...")
    from thesis_compiler.chapters.preliminary import METADATA
    
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(36)
    title_p.paragraph_format.space_after = Pt(24)
    run_title = title_p.add_run(f"{METADATA['title']}\n\n")
    apply_text_formatting(title_p, font_name="Georgia", size_pt=18, bold=True, color_rgb=(27, 54, 93))
    
    body_p = doc.add_paragraph()
    body_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    body_p.paragraph_format.space_before = Pt(48)
    body_p.paragraph_format.space_after = Pt(12)
    body_p.add_run("A dissertation manuscript submitted in partial fulfillment of the requirements for the degree of\n")
    body_p.add_run(f"{METADATA['degree']}\n\n")
    apply_text_formatting(body_p, size_pt=11.5, italic=True)
    
    cand_p = doc.add_paragraph()
    cand_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cand_p.paragraph_format.space_before = Pt(48)
    cand_p.add_run("Submitted by\n")
    run_cand = cand_p.add_run(f"{METADATA['candidate']}\n\n")
    apply_text_formatting(cand_p, size_pt=12)
    run_cand.bold = True
    
    super_p = doc.add_paragraph()
    super_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    super_p.add_run("Under the Supervision of\n")
    run_super = super_p.add_run(f"{METADATA['supervisor']}\n")
    super_p.add_run(f"{METADATA['supervisor_designation']}\n")
    super_p.add_run(f"{METADATA['department']}\n")
    super_p.add_run(f"{METADATA['institution']}\n")
    apply_text_formatting(super_p, size_pt=11)
    run_super.bold = True
    
    univ_p = doc.add_paragraph()
    univ_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    univ_p.paragraph_format.space_before = Pt(64)
    univ_p.add_run(f"{METADATA['institution']}\n")
    univ_p.add_run(f"{METADATA['university']}\n")
    univ_p.add_run(f"Academic Year: {METADATA['year']}\n")
    apply_text_formatting(univ_p, size_pt=11.5, bold=True, color_rgb=(27, 54, 93))
    
    # ─── Preliminary Pages ────────────────────────────────────────────────────
    print("Compiling Preliminary Front Matter...")
    from thesis_compiler.chapters import preliminary
    
    for item in preliminary.CONTENT:
        tag = item[0]
        if tag == "page_break":
            doc.add_page_break()
        elif tag == "h1":
            p = doc.add_heading(item[1], level=1)
            p.paragraph_format.space_before = Pt(24)
            p.paragraph_format.space_after = Pt(12)
            apply_text_formatting(p, font_name="Georgia", size_pt=16, bold=True, color_rgb=(27, 54, 93))
            p.paragraph_format.keep_with_next = True
        elif tag == "p":
            p = doc.add_paragraph(item[1])
            apply_text_formatting(p)
        elif tag == "table":
            headers, rows, caption = item[1], item[2], item[3]
            table = doc.add_table(rows=len(rows)+1, cols=len(headers))
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            set_table_borders(table)
            
            # Header Row
            hdr_cells = table.rows[0].cells
            for idx, header_text in enumerate(headers):
                hdr_cells[idx].text = header_text
                set_cell_background(hdr_cells[idx], "1B365D")
                p = hdr_cells[idx].paragraphs[0]
                apply_text_formatting(p, size_pt=9.5, bold=True, color_rgb=(255, 255, 255))
            
            # Data Rows
            for r_idx, row_data in enumerate(rows):
                row_cells = table.rows[r_idx+1].cells
                bg_color = "F8FAFC" if r_idx % 2 == 0 else "FFFFFF"
                for c_idx, val in enumerate(row_data):
                    row_cells[c_idx].text = str(val)
                    set_cell_background(row_cells[c_idx], bg_color)
                    p = row_cells[c_idx].paragraphs[0]
                    apply_text_formatting(p, size_pt=9)
            
            cap_p = doc.add_paragraph()
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_p.add_run(caption).italic = True
            apply_text_formatting(cap_p, size_pt=8.5, color_rgb=(100, 100, 100))
            
    # Add Table of Contents Dynamic Field
    doc.add_page_break()
    toc_h = doc.add_heading("Table of Contents", level=1)
    apply_text_formatting(toc_h, font_name="Georgia", size_pt=16, bold=True, color_rgb=(27, 54, 93))
    toc_h.paragraph_format.keep_with_next = True
    
    toc_p = doc.add_paragraph()
    toc_p.paragraph_format.space_before = Pt(12)
    add_toc(toc_p)
    apply_text_formatting(toc_p, size_pt=11)
    
    # ─── Chapter-by-Chapter Compiling ─────────────────────────────────────────
    print("Compiling Thesis Body Chapters...")
    for chap in chapters_list:
        print(f"  Processing: {chap.title}")
        doc.add_page_break()
        
        for item in chap.CONTENT:
            tag = item[0]
            if tag == "h1":
                p = doc.add_heading(item[1], level=1)
                p.paragraph_format.space_before = Pt(36)
                p.paragraph_format.space_after = Pt(12)
                apply_text_formatting(p, font_name="Georgia", size_pt=16, bold=True, color_rgb=(27, 54, 93))
                p.paragraph_format.keep_with_next = True
            elif tag == "h2":
                p = doc.add_heading(item[1], level=2)
                p.paragraph_format.space_before = Pt(24)
                p.paragraph_format.space_after = Pt(8)
                apply_text_formatting(p, font_name="Georgia", size_pt=13, bold=True, color_rgb=(27, 54, 93))
                p.paragraph_format.keep_with_next = True
            elif tag == "h3":
                p = doc.add_heading(item[1], level=3)
                p.paragraph_format.space_before = Pt(16)
                p.paragraph_format.space_after = Pt(6)
                apply_text_formatting(p, font_name="Georgia", size_pt=11.5, bold=True, color_rgb=(71, 85, 105))
                p.paragraph_format.keep_with_next = True
            elif tag == "p":
                p = doc.add_paragraph(item[1])
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                apply_text_formatting(p)
            elif tag == "equation":
                eq_p = doc.add_paragraph()
                eq_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                eq_p.paragraph_format.space_before = Pt(12)
                eq_p.paragraph_format.space_after = Pt(12)
                eq_p.add_run(f"$$ {item[1]} $$")
                apply_text_formatting(eq_p, size_pt=11.5, italic=True)
            elif tag == "list":
                for bullet in item[1]:
                    p = doc.add_paragraph(style='List Bullet')
                    p.paragraph_format.space_after = Pt(3)
                    p.add_run(bullet)
                    apply_text_formatting(p)
            elif tag == "figure":
                fig_name, caption = item[1], item[2]
                fig_path = os.path.join(figures_dir, fig_name)
                if os.path.exists(fig_path):
                    fig_p = doc.add_paragraph()
                    fig_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    fig_p.paragraph_format.space_before = Pt(12)
                    fig_p.paragraph_format.space_after = Pt(6)
                    fig_p.add_run().add_picture(fig_path, width=Inches(5.5))
                    
                    cap_p = doc.add_paragraph()
                    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap_p.add_run(caption).italic = True
                    apply_text_formatting(cap_p, size_pt=8.5, color_rgb=(100, 100, 100))
                else:
                    print(f"  [WARNING] Figure not found: {fig_path}")
            elif tag == "table":
                headers, rows, caption = item[1], item[2], item[3]
                table = doc.add_table(rows=len(rows)+1, cols=len(headers))
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                set_table_borders(table)
                
                # Header Row
                hdr_cells = table.rows[0].cells
                for idx, header_text in enumerate(headers):
                    hdr_cells[idx].text = header_text
                    set_cell_background(hdr_cells[idx], "1B365D")
                    p = hdr_cells[idx].paragraphs[0]
                    apply_text_formatting(p, size_pt=9.5, bold=True, color_rgb=(255, 255, 255))
                
                # Data Rows
                for r_idx, row_data in enumerate(rows):
                    row_cells = table.rows[r_idx+1].cells
                    bg_color = "F8FAFC" if r_idx % 2 == 0 else "FFFFFF"
                    for c_idx, val in enumerate(row_data):
                        row_cells[c_idx].text = str(val)
                        set_cell_background(row_cells[c_idx], bg_color)
                        p = row_cells[c_idx].paragraphs[0]
                        apply_text_formatting(p, size_pt=9)
                
                cap_p = doc.add_paragraph()
                cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap_p.add_run(caption).italic = True
                apply_text_formatting(cap_p, size_pt=8.5, color_rgb=(100, 100, 100))
            elif tag == "code":
                code_text, caption = item[1], item[2]
                code_table = doc.add_table(rows=1, cols=1)
                code_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                cell = code_table.rows[0].cells[0]
                set_cell_background(cell, "F1F5F9")
                
                # Add border to table cell via XML
                tcPr = cell._tc.get_or_add_tcPr()
                borders_xml = f'''<w:tcBorders {nsdecls("w")}>
                    <w:left w:val="single" w:sz="18" w:space="0" w:color="1B365D"/>
                    <w:top w:val="none"/>
                    <w:right w:val="none"/>
                    <w:bottom w:val="none"/>
                </w:tcBorders>'''
                tcPr.append(parse_xml(borders_xml))
                
                cell_p = cell.paragraphs[0]
                cell_p.text = code_text
                apply_text_formatting(cell_p, font_name="Consolas", size_pt=8.5, color_rgb=(15, 23, 42))
                cell_p.paragraph_format.line_spacing = 1.0
                cell_p.paragraph_format.space_before = Pt(4)
                cell_p.paragraph_format.space_after = Pt(4)
                
                cap_p = doc.add_paragraph()
                cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap_p.add_run(caption).italic = True
                apply_text_formatting(cap_p, size_pt=8.5, color_rgb=(100, 100, 100))
            elif tag == "page_break":
                doc.add_page_break()

    # ─── Bibliography Page ────────────────────────────────────────────────────
    print("Compiling Reference Bibliography Page (Vancouver style)...")
    doc.add_page_break()
    ref_h = doc.add_heading("References", level=1)
    apply_text_formatting(ref_h, font_name="Georgia", size_pt=16, bold=True, color_rgb=(27, 54, 93))
    ref_h.paragraph_format.keep_with_next = True
    
    # Sort references by Vancouver number (appearance order)
    sorted_ref_keys = sorted(citation_map.keys(), key=lambda k: citation_map[k])
    
    for key in sorted_ref_keys:
        num = citation_map[key]
        ref_text = REFERENCES[key]
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.space_after = Pt(8)
        p.add_run(f"[{num}] {ref_text}")
        apply_text_formatting(p, size_pt=10.5)

    # ─── Save Document ────────────────────────────────────────────────────────
    print(f"Writing Word output file to {output_path}...")
    doc.save(output_path)
    print("Word Compilation successfully completed!")

if __name__ == "__main__":
    from thesis_compiler.chapters import CHAPTERS
    figures_directory = os.path.join(os.path.dirname(__file__), "figures")
    output_filename = os.path.join(os.path.dirname(__file__), "thesis_manuscript.docx")
    build_docx(output_filename, figures_directory, CHAPTERS)
